import os
from docx import Document
from docx.shared import Inches

def replace_header_assets(doc, logo_path, doc_prefix="DG"):
    """Replaces logo in header and updates Doc No prefix."""
    for section in doc.sections:
        header = section.header
        
        # 1. Replace Logo (Assuming it's in a table cell like previous versions)
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "MHS" in cell.text or "Doc No" in cell.text:
                        # Logic to update Doc No prefix (e.g., ABC-MHS -> DG-MHS)
                        for paragraph in cell.paragraphs:
                            if "Doc No" in paragraph.text:
                                # Regex-style replacement of the first part before the first hyphen
                                parts = paragraph.text.split('-')
                                if len(parts) > 1:
                                    # Keeps the rest of the ID, replaces the first segment
                                    paragraph.text = f"Doc No: {doc_prefix}-" + "-".join(parts[1:])
                    
                    # Remove old images and add new logo
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"):
                                run.clear()
                                run.add_picture(logo_path, width=Inches(1.5))

def replace_footer_assets(doc, footer_path):
    """Replaces all footer content with the new footer image."""
    for section in doc.sections:
        footer = section.footer
        for paragraph in footer.paragraphs:
            paragraph.clear() # Remove old text/images
        
        # Add the new footer image to the first paragraph
        footer.paragraphs[0].add_run().add_picture(footer_path, width=Inches(6.0))

def replace_body_text(doc):
    """Handles specific text replacements in the document body."""
    # Mapping of {Old Text : New Text}
    replacements = {
        "Souring Summits Group Pty (Ltd)": "Dynatron Group (Pty) Ltd",
        "Scope of Work: The provision of engineering consulting, project and maintenance management.": 
        "Scope of Works: Provision of Engineering, Construction and Project Management Services"
    }

    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                paragraph.text = paragraph.text.replace(old, new)
                # Note: To keep bolding, we'd need run-level processing, 
                # but standard .replace is safer for long sentences.

def process_system():
    base_path = r'C:\Users\Admin\Documents\Work\Mntambo Services\Document Processing\Clients\01'
    logo_path = os.path.join(base_path, 'assets', 'logo', '01.png')
    footer_path = os.path.join(base_path, 'assets', 'footer', '01.png')
    output_dir = os.path.join(base_path, 'outcome')

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for filename in os.listdir(base_path):
        if filename.endswith(".docx") and not filename.startswith("~$"):
            file_path = os.path.join(base_path, filename)
            print(f"Processing: {filename}...")
            
            try:
                doc = Document(file_path)
                
                replace_header_assets(doc, logo_path)
                replace_footer_assets(doc, footer_path)
                replace_body_text(doc)
                
                doc.save(os.path.join(output_dir, filename))
            except Exception as e:
                print(f"Failed to process {filename}: {e}")

if __name__ == "__main__":
    process_system()